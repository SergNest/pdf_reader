import base64
import logging
import os
import re
from datetime import datetime
from io import BytesIO
from typing import List, Tuple

from flask import Flask, request, render_template, send_file, redirect, url_for, abort, jsonify
from werkzeug.utils import secure_filename
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from dotenv import load_dotenv
from PIL import Image, ImageFilter
from werkzeug.exceptions import RequestEntityTooLarge
import openai
from openai import OpenAI

logging.basicConfig(level=logging.INFO)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 1 MB
app.config['UPLOAD_FOLDER'] = os.getenv("UPLOAD_FOLDER", "uploads")
app.config['CONVERTED_FOLDER'] = os.getenv("CONVERTED_FOLDER", "converted_files")

allowed_ips = [ip.strip() for ip in os.getenv('MY_LIST', '').split(',') if ip.strip()]
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg'}

client = OpenAI(
    api_key=os.getenv("YOUR_API_KEY", "api_key"),
)


# Перевірка розширення файлу
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.before_request
def limit_remote_addr():
    if request.remote_addr not in allowed_ips:
        abort(403)  # Доступ заборонено


def convert_to_docx(file_path, output_filename, use_ai=False):
    """
    Converts a PDF or image file to a DOCX file.

    Parameters:
    file_path (str): The path to the input file (PDF or image).
    output_filename (str): The name of the output DOCX file.
    use_ai (bool, optional): If True, uses OpenAI API to correct the extracted text. Defaults to False.

    Returns:
    str: The path to the output DOCX file.
    """

    doc = Document()
    _, ext = os.path.splitext(file_path)
    custom_oem_psm_config = r'--oem 3 --psm 6 -l eng+best'

    if ext.lower() == '.pdf':
        images = convert_from_path(file_path)
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image, config=custom_oem_psm_config)
            text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\xA0-\uD7FF\uE000-\uFFFD]', '', text)
            if use_ai:
                text = correct_text_with_ai(text)
            doc.add_paragraph(text)

    elif ext.lower() in {'.png', '.jpg', '.jpeg'}:
        image = Image.open(file_path)
        image = image.convert('L')  # Перетворення в градації сірого
        image = image.filter(ImageFilter.SHARPEN)  # Застосування фільтра підвищення різкості
        text = pytesseract.image_to_string(image, config=custom_oem_psm_config)
        text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\xA0-\uD7FF\uE000-\uFFFD]', '', text)
        if use_ai:
            text = correct_text_with_ai(text)
        doc.add_paragraph(text)

    output_path = str(os.path.join(app.config['CONVERTED_FOLDER'], output_filename))
    doc.save(output_path)
    return output_path


@app.route("/", methods=["GET", "POST"])
def index():
    try:
        if request.method == 'POST':

            file = request.files.get('file')
            pasted_image = request.form.get('pasted_image')

            # Process pasted image
            if pasted_image:
                image_data = base64.b64decode(pasted_image.split(',')[1])
                image = Image.open(BytesIO(image_data))
                filename = "pasted_image.png"
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                image.save(file_path)
                output_filename = f"pasted_image_{datetime.now().strftime('%d%H%M')}.docx"
                use_ai = request.form.get('use_ai') == 'on'
                convert_to_docx(file_path, output_filename, use_ai=use_ai)

            # Process uploaded file
            elif file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                output_filename = f"{os.path.splitext(filename)[0]}.docx"
                use_ai = request.form.get('use_ai') == 'on'
                convert_to_docx(file_path, output_filename, use_ai=use_ai)

            else:
                return 'No file or pasted image selected'

    except RequestEntityTooLarge as e:
        app.logger.error(f"Request entity too large: {e}")
        return "Request entity too large. Please try uploading a smaller file.", 413

    converted_files = get_last_converted_files()
    return render_template('index.html', converted_files=converted_files)


@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    # Генерація імені файлу
    filename = "pasted_image1.png"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    # Збереження завантаженого файлу
    file.save(file_path)

    # Генерація імені вихідного файлу
    output_filename = f"pasted_image_{datetime.now().strftime('%d%H%M')}.docx"
    use_ai = request.form.get('use_ai') == 'true'

    # Конвертація зображення в DOCX
    convert_to_docx(file_path, output_filename, use_ai=use_ai)

    return jsonify({'filename': output_filename}), 200


@app.route("/download/<filename>")
def download_file(filename):
    path = os.path.join(app.config['CONVERTED_FOLDER'], filename)
    return send_file(path, as_attachment=True)


@app.route('/last_files', methods=['GET'])
def last_files():
    files = get_last_converted_files()
    return jsonify({'files': files})

def get_last_converted_files(converted_folder: str = app.config['CONVERTED_FOLDER'], num_files: int = 5) -> List[str]:
    """
    Retrieves the names of the last converted files in the specified converted folder.

    Parameters:
    - converted_folder (str): The path to the folder containing the converted files. Defaults to the value of the 'CONVERTED_FOLDER' configuration variable.
    - num_files (int): The number of files to retrieve. Defaults to 5.

    Returns:
    - List[str]: A list of the names of the last converted files, sorted by modification time in descending order.

    This function retrieves the names of the last converted files in the specified converted folder. It uses os.scandir for better performance and retrieves the modification time of each file. The files are then sorted by modification time in descending order, and the first 'num_files' files are returned.
    """
    if not os.path.exists(converted_folder):
        return []

    files_with_time: List[Tuple[str, float]] = []
    for entry in os.scandir(converted_folder):
        if entry.is_file():
            try:
                mtime = entry.stat().st_mtime
                files_with_time.append((entry.name, mtime))
            except OSError as e:
                print(f"Помилка при отриманні часу модифікації для файлу {entry.path}: {e}")
                continue  # Пропускаємо файл, якщо виникла помилка

    # Сортуємо за часом модифікації, від найновіших до найстаріших
    files_with_time.sort(key=lambda x: x[1], reverse=True)

    return [f[0] for f in files_with_time[:num_files]]


def correct_text_with_ai(text: str) -> str:
    """
    Uses the OpenAI GPT-3.5 model to correct the extracted text.

    Parameters:
    text (str): The extracted text to be corrected.

    Returns:
    str: The corrected text.

    This function sends a request to the OpenAI API using the GPT-3.5 model to correct the extracted text. It constructs a message containing the original text and sends it to the model. The model then generates a corrected version of the text, which is returned. If there is an error connecting to the OpenAI API, the original text is returned.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a text corrector."},
                {"role": "user", "content": f"Correct the following text:\n\n{text}"}
            ],
            max_tokens=1000
        )
        return response.choices[0].message.content.strip()
    except openai.APIConnectionError as e:
        print(f"OpenAI API Connection Error: {e}")
        return text


if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)
    app.run(debug=True, port=5006, host='0.0.0.0')






